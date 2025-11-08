"""Export manager for generating PDF, CSV, and DOCX reports.
Handles optional dependencies gracefully and provides simple helpers
for exporting lists of dicts.
"""
from __future__ import annotations

from typing import List, Dict, Any
from datetime import datetime

# Optional deps
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    _HAS_PDF = True
except Exception:
    _HAS_PDF = False

try:
    import pandas as pd  # type: ignore
    _HAS_CSV = True
except Exception:
    _HAS_CSV = False

try:
    from docx import Document  # type: ignore
    from docx.shared import Pt
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False


def export_to_pdf(path: str, title: str, records: List[Dict[str, Any]]):
    if not _HAS_PDF:
        raise ImportError("reportlab is not installed")
    doc = SimpleDocTemplate(path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    header = styles['Title']
    story.append(Paragraph("SENTRY REPORT", header))
    story.append(Spacer(1, 12))
    sub = styles['Normal']
    story.append(Paragraph(f"{title}", sub))
    story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", sub))
    story.append(Spacer(1, 12))

    if not records:
        story.append(Paragraph("Sem dados.", styles['Italic']))
        doc.build(story)
        return

    # Build a table from keys/values
    # Normalize keys across records
    cols = list({k for r in records for k in r.keys()})
    data = [cols]
    for r in records:
        row = [str(r.get(c, '')) for c in cols]
        data.append(row)

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.black),
        ('GRID', (0,0), (-1,-1), 0.25, colors.grey),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.whitesmoke, colors.beige])
    ]))

    story.append(table)
    doc.build(story)


def export_to_csv(path: str, records: List[Dict[str, Any]]):
    if not _HAS_CSV:
        raise ImportError("pandas is not installed")
    import pandas as pd  # type: ignore
    df = pd.DataFrame(records or [])
    df.to_csv(path, index=False, encoding='utf-8')


def export_to_docx(path: str, title: str, records: List[Dict[str, Any]]):
    if not _HAS_DOCX:
        raise ImportError("python-docx is not installed")
    doc = Document()
    doc.add_heading('SENTRY REPORT', level=0)
    doc.add_paragraph(title)
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(" ")

    if not records:
        doc.add_paragraph('Sem dados.')
        doc.save(path)
        return

    # Use a table with columns from union of keys
    cols = list({k for r in records for k in r.keys()})
    table = doc.add_table(rows=1, cols=len(cols))
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr_cells[i].text = str(c)

    for r in records:
        row_cells = table.add_row().cells
        for i, c in enumerate(cols):
            row_cells[i].text = str(r.get(c, ''))

    doc.save(path)


def get_dependency_status() -> Dict[str, bool]:
    return {
        'pdf': _HAS_PDF,
        'csv': _HAS_CSV,
        'docx': _HAS_DOCX,
    }
